from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/bytedance/Documents/Codex/new2/汇报材料")
REPORT_PATH = ROOT / "KV智能生图工具项目汇报.docx"
SPEECH_PATH = ROOT / "KV智能生图工具5分钟演讲稿.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183B56"
TEAL = "18A999"
PALE_BLUE = "EAF3F8"
PALE_TEAL = "E8F7F4"
PALE_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK = "18212B"
WHITE = "FFFFFF"
ORANGE = "F2A65A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D9DEE5", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeatable_header_footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        header = section.header
        hp = header.paragraphs[0]
        hp.text = title
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style = doc.styles["Caption"]
        for run in hp.runs:
            run.font.color.rgb = RGBColor.from_string(MID_GRAY)
            run.font.size = Pt(8)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("第 ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        fp._p.append(fld)
        fp.add_run(" 页")
        for fr in fp.runs:
            fr.font.color.rgb = RGBColor.from_string(MID_GRAY)
            fr.font.size = Pt(8)


def set_font(run, name="Hiragino Sans GB", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {
        "Title": (28, NAVY, 0, 14),
        "Subtitle": (12, MID_GRAY, 0, 8),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Hiragino Sans GB"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Hiragino Sans GB"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    caption.font.size = Pt(8)
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Hiragino Sans GB"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def add_cover(doc: Document, title: str, subtitle: str, kicker: str, note: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(kicker.upper())
    set_font(r, size=9, bold=True, color=TEAL)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_font(r, size=13, color=MID_GRAY)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    set_cell_margins(cell, top=220, start=240, bottom=220, end=240)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(note)
    set_font(r, size=13, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(180)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("项目阶段：可运行原型 / 生产级策略迭代中")
    set_font(r, size=9, color=MID_GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("汇报日期：2026 年 7 月")
    set_font(r, size=9, color=MID_GRAY)
    doc.add_page_break()


def add_callout(doc: Document, title: str, text: str, color=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)


def add_number(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE_GRAY)
        set_cell_margins(cell)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_pipeline(doc: Document) -> None:
    steps = [
        ("01", "输入", "标题 / 副标题 / 时间 / 画面描述 / 比例 / 上传图"),
        ("02", "Brief", "只提取明确事实，区分硬约束与推断"),
        ("03", "创意判断", "3 个方向脑暴，自评淘汰 2 个，选定 1 个"),
        ("04", "参考匹配", "整合版式 / 风格 / 元素 / 角色，记录用途与原因"),
        ("05", "设计大纲", "主体、占比、层级、色彩、材质、光影、镜头"),
        ("06", "美术预检", "覆盖度、创意强度、可执行性、留白与比例"),
        ("07", "Prompt + 图", "结构化 Prompt 与真实参考图一起提交"),
        ("08", "成图与资产", "生成、叠加固定层、存档、拆分与复用"),
    ]
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    idx = 0
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            number, title, desc = steps[idx]
            idx += 1
            cell.width = Inches(3.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cell, PALE_TEAL if idx % 2 else PALE_BLUE)
            set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
            set_cell_border(cell, color="FFFFFF", size="12")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{number}  {title}")
            set_font(r, size=10.5, bold=True, color=DARK_BLUE)
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(desc)
            set_font(r, size=8.8, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_report() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(
        doc,
        "营销 KV 智能生图工具",
        "项目汇报：从一句需求，到有设计判断、有参考约束、可沉淀复用的生产链路",
        "PROJECT REPORT",
        "它不是一个“输入一句话、碰运气出图”的页面，而是一套把设计师工作拆成可执行步骤，并让真实参考图进入生成接口的营销视觉工作台。",
    )
    set_repeatable_header_footer(doc, "营销 KV 智能生图工具｜项目汇报")

    doc.add_heading("汇报摘要", level=1)
    add_callout(
        doc,
        "一句话结论",
        "当前版本已完成“需求输入 → 设计理解 → 创意判断 → 参考匹配 → 结构化 Prompt → 参考图随 Prompt 进入模型 → 成图 → 资产沉淀/拆分/复用”的主链路，并进一步跑通了风格与素材持续扩充、外部灵感搜索保存、用户上传图引用等多个闭环。",
    )
    add_table(
        doc,
        ["要回答的问题", "当前答案"],
        [
            ["为什么做", "营销需求通常很简短，但高质量 KV 需要创意、版式、风格、素材和信息层级共同成立。只写 Prompt 很难稳定复现这套判断。"],
            ["核心方法", "把设计流程拆成多个可检查节点，并把参考图按用途真实传入图像模型，而不是只在文字里描述“参考某图”。"],
            ["目前做到什么", "可运行生成工作台、4 个启用风格预设 + 默认模式、整合版式、上传图与 @ 引用、兜兜 IP、素材/风格/资产管理、Pinterest 与 Behance 搜索等。"],
            ["形成什么闭环", "需求到成图、灵感到素材、素材到生成、生成到资产、资产到分层、风格文件夹到新预设均已首尾相接。"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("一、背景：为什么要做这个工具", level=1)
    doc.add_paragraph(
        "营销 KV 的难点并不只是“把画画出来”。真正决定成片质量的，是对业务信息的取舍、创意概念、主体关系、镜头构图、文字层级、参考图用途以及最终执行约束。业务和运营同学往往只能提供一句简短需求，设计师则需要把这句话补成一套完整方案。"
    )
    doc.add_heading("1. 当前工作里的四个痛点", level=2)
    add_bullet(doc, "输入天然不完整：用户经常只给活动名称和一句画面描述，但模型容易擅自补人物、品牌、英文、卖点或无关道具。")
    add_bullet(doc, "一步生图不可控：Prompt 写得再长，也不等于模型真的理解了字体、版式和参考图各自该控制什么。")
    add_bullet(doc, "经验难以复用：优秀案例、风格图、元素图和设计方法散落在个人文件夹里，下一次仍要从头寻找。")
    add_bullet(doc, "结果难以继续改：一步出图对非设计师方便，但设计师接手后往往缺少可拆分、可替换、可复用的中间资产。")
    doc.add_heading("2. 工具的目标", level=2)
    add_bullet(doc, "降低输入门槛：用户不需要会写专业 Prompt，也能得到完整设计方案。")
    add_bullet(doc, "增加设计思考：在出图前加入创意发散、方案筛选和美术预检。")
    add_bullet(doc, "提高参考可控性：每张参考图明确用途，并作为真实图片提交给生成接口。")
    add_bullet(doc, "让内容可复用：素材、风格、案例、生成结果和拆分结果都能回到系统继续使用。")
    add_bullet(doc, "平衡速度与质量：默认快速模式控制等待时间，质量模式保留更完整的模型评审与重试。")

    doc.add_heading("二、产品定位与使用对象", level=1)
    add_callout(
        doc,
        "产品定位",
        "面向营销活动视觉的“设计推理 + 参考图约束 + 图像生成 + 资产沉淀”一体化工作台。它不是要替代设计师，而是把原本只存在于设计师脑中的步骤显性化，让业务和运营能得到更接近可用稿的第一版，也让设计师接手时拥有更清楚的依据。",
        PALE_BLUE,
    )
    add_table(
        doc,
        ["用户", "主要诉求", "工具提供的价值"],
        [
            ["业务 / 运营", "输入少、时间紧，希望快速获得可用 KV", "自然语言输入、风格选择、扩写、流式进度、直接成图"],
            ["设计师", "需要更可控、更容易继续修改", "参考图分工、设计大纲、Prompt 展开、资产拆分、素材复用"],
            ["设计管理者", "希望沉淀方法和标准，而不是只沉淀成图", "创意方法卡、Good/Bad Case、美术预检、预设知识库"],
        ],
        [1.25, 2.2, 3.05],
    )

    doc.add_heading("三、端到端生图链路", level=1)
    add_pipeline(doc)
    doc.add_paragraph(
        "这条链路的关键不是节点多，而是每个节点都承担不同职责：Brief 不负责想象画面，创意判断不直接写最终 Prompt，参考匹配不随机抽图，最终生图也不能脱离真实参考图单独调用。"
    )

    doc.add_heading("四、生图策略与算法逻辑", level=1)
    doc.add_heading("1. 输入解析：先锁定事实，再允许推断", level=2)
    doc.add_paragraph(
        "系统读取活动名称、可选副标题、可选活动时间、画面描述、画幅比例、风格预设、上传图片及其 @ 用途，以及 Logo / 搜索框 / 兜兜 IP 等开关。Brief 节点只提取用户明确给出的事实，并将“硬约束”和“合理推断”分开，避免把模型的想象误当成用户需求。"
    )
    add_bullet(doc, "硬约束示例：主标题文字、画幅 3:4、上传图是唯一产品主体、禁止出现额外英文。")
    add_bullet(doc, "可推断项示例：用户写“清爽夏日”，可推断低负担、明亮、留白，但不能自行添加折扣、价格或品牌。")

    doc.add_heading("2. 创意判断：先想方案，再选方案", level=2)
    doc.add_paragraph(
        "系统读取当前风格的 preset.md、12 张创意方法卡和 Good/Bad Case。它必须先提出 3 个差异化方向，再从主题覆盖、记忆点、画幅适配、执行风险等角度自行评估，淘汰 2 个并选择 1 个。最终保留的不只是“风格词”，还包括视觉载体、冲突趣味、记忆符号和分层草图。"
    )
    add_callout(
        doc,
        "示例：只输入“冰雪温泉季竖版海报”",
        "方案 1 只表现滑雪，缺失温泉；方案 2 左右各一半，竖版阅读割裂；方案 3 采用纵深雪景，把滑雪角色、温泉角色、远景小镇和顶部标题整合在同一场景中。系统选择方案 3，再据此匹配参考图。",
        PALE_TEAL,
    )

    doc.add_heading("3. 参考图检索：从随机抽取改为“先判断用途，再匹配证据”", level=2)
    doc.add_paragraph(
        "参考图按职责拆分。当前主要包含整合版式、风格质感、元素主体和角色等类型。检索不是简单随机，而是先生成检索意图，再进行本地规则筛选与语义评分；质量模式可对候选集增加 LLM 重排。每张入选图片都会记录“为什么选、用于参考什么、哪些内容禁止复制”。"
    )
    add_table(
        doc,
        ["参考类型", "主要控制内容", "不应控制内容"],
        [
            ["整合版式", "标题、副标题、时间的样式与比例；信息区、主画面区、留白区；对齐与阅读顺序；原图已有装饰", "本次画面的具体主体和最终颜色；原参考中的业务文案、品牌、版权"],
            ["风格质感", "整体美术语言、完成度、材质、光影气质", "标题字形、具体产品、参考图原文案"],
            ["元素主体", "造型概括、比例、边缘、材质处理", "元素原品牌、原颜色和无关道具"],
            ["角色", "角色比例、轮廓、动作重心、运动方向", "角色身份、原服装文字、品牌和背景"],
            ["用户上传图", "产品或人物身份、结构、轮廓、包装比例和可见识别特征", "将产品替换成相似品类或无关主体"],
        ],
        [1.15, 3.05, 2.3],
    )

    doc.add_heading("4. 整合版式：文字视觉系统和 KV 区域关系由同一张图约束", level=2)
    doc.add_paragraph(
        "当前版式策略已从“字体图 + 日期图 + 排版图分别控制”升级为整合版式。系统根据画幅方向、标题字数与行数、是否存在副标题和时间、信息槽位等条件选择一张整合版式参考，并用它同时控制文字样式和画面区域。"
    )
    add_bullet(doc, "参考图中已有的引号、框线、标签、角标、下划线、色块和装饰符号需要保留。")
    add_bullet(doc, "参考图中已有的英文眉题等非业务装饰文字原样保留，不改写；参考图没有的内容不额外发散。")
    add_bullet(doc, "主标题、副标题和时间替换为用户输入；原参考的品牌、版权和业务文案删除或替换。")
    add_bullet(doc, "白色或空白区域表示主视觉生成区，不代表最终背景必须是白色。")
    add_bullet(doc, "文字颜色不跟随版式参考图，而根据本次画面重选；副标题和时间与主标题保持同一颜色系统。")
    add_bullet(doc, "对齐方式依据参考图和本次字数决定，不默认全部左对齐。")

    doc.add_heading("5. 设计大纲：把创意变成可执行参数", level=2)
    doc.add_paragraph(
        "选定创意和参考后，系统生成设计大纲，定义整体视觉关键词、色彩方向、主视觉主体、主体占比、主体与其他元素关系、信息层级、背景、材质、光影、镜头及构图。它相当于生图前的“画面施工图”，后续 Prompt 只能在这份大纲范围内展开。"
    )

    doc.add_heading("6. 美术预检：出图前先做一次总监式检查", level=2)
    doc.add_paragraph(
        "预检覆盖 Brief 是否完整、创意是否有记忆点、主体是否明确、版式是否适配比例、参考证据是否充分、留白和信息层级是否可执行。发现问题时只做最小修补，避免整套方案被二次发散。质量模式还可启用生成后评审，并在必要时做一次有方向的重试。"
    )

    doc.add_heading("7. Prompt 拼装与生成：真实参考图必须随请求提交", level=2)
    doc.add_paragraph(
        "最终 Prompt 采用结构化章节，只写有来源的信息，不为缺失字段擅自补内容。系统随后把 Prompt 与已选整合版式、风格、元素、角色、兜兜或用户上传图等真实图片一起提交给 gpt-image-2。参考关系因此不是“文字里说参考了”，而是模型实际收到图片。"
    )
    add_bullet(doc, "参考优先级清楚：用户指定主体身份 > 整合版式的文字/区域关系 > 风格 > 元素/角色。")
    add_bullet(doc, "用户上传图作为主体时，必须保留品类、结构、轮廓和包装识别，禁止把洗衣液生成成手机之类的主体漂移。")
    add_bullet(doc, "选择兜兜 IP 后，系统按设计判断匹配多张姿态参考；兜兜必须出现，并遵守无手、无胳膊、无嘴等角色约束。")
    add_bullet(doc, "Logo 和右下角活动搜索框为可选后处理图层，默认关闭；启用后使用真实素材叠加，不让模型重绘。")

    doc.add_heading("8. 速度策略：快速模式与质量模式", level=2)
    doc.add_paragraph(
        "默认快速模式将 Brief、基础设计、检索和预检尽量放在本地执行，把网络调用集中在创意规划、Prompt 优化和最终生图，目标控制在 3 分钟内。质量模式可以打开更多 LLM 节点、生成后评审和一次自动重试。页面通过 SSE 流式返回 Brief、设计判断、Prompt 和生成状态，避免长时间无反馈。"
    )

    doc.add_heading("五、目前已经实现的能力", level=1)
    add_table(
        doc,
        ["模块", "已实现能力", "状态"],
        [
            ["生成工作台", "标题 / 副标题 / 时间 / 描述 / 比例输入；多图上传；@ 引用；扩写；风格、兜兜、Logo、搜索框开关；流式节点展示", "已可用"],
            ["风格预设", "默认无预设 + 手绘扁平涂鸦、3D、极简扁平插画、实景商品；支持文件夹导入、配置、删除", "已可用"],
            ["参考系统", "整合版式、风格、元素、角色、上传主体与兜兜多参考；记录用途与选择原因；图片真实传入生成接口", "已可用，匹配精度持续优化"],
            ["创意系统", "12 张创意方法卡；Good/Bad Case；3 方案发散、淘汰与选定；美术预检", "已接入主链路"],
            ["素材库", "7 类筛选；上传、XLSX 导入、删除、详情查看；用作参考图；做同款；预设素材同步展示", "已可用"],
            ["灵感搜索", "Pinterest + Behance 多关键词并行搜索、过滤、去重、来源回退、预览、保存到素材库", "验证可用"],
            ["资产页", "生成结果倒序展示、删除、查看原始输入；AI 拆分标题层与背景层；下载拆分包", "已可用"],
            ["固定图层", "浅/深色 Logo 自动选择；浅/深色活动搜索框自动选择；搜索框写入活动名称", "已可用"],
        ],
        [1.1, 4.55, 0.85],
    )

    doc.add_heading("六、已经跑通的业务闭环", level=1)
    add_number(doc, "需求到成图闭环：自然语言需求进入系统，经过 Brief、创意、参考、设计、预检和 Prompt，最终携带真实参考图完成生图。", "需求到成图闭环：")
    add_number(doc, "灵感到素材闭环：从 Pinterest / Behance 搜索案例，选中后保存到素材库，之后可继续作为参考图参与生成。", "灵感到素材闭环：")
    add_number(doc, "素材到生成闭环：素材详情页可直接“用作参考图”回到当前生成表单，也可“做同款”把描述带入画面输入。", "素材到生成闭环：")
    add_number(doc, "风格扩充闭环：按约定文件夹结构上传新风格，系统读取图片和描述形成预设，并在首页供用户选择。", "风格扩充闭环：")
    add_number(doc, "生成到资产闭环：生成图自动进入资产页，可查看、删除、继续下载和拆分。", "生成到资产闭环：")
    add_number(doc, "资产到可编辑闭环：AI 可从成图中分离标题文字层和清理后的背景层，形成可继续加工的拆分包。", "资产到可编辑闭环：")
    add_number(doc, "IP 使用闭环：兜兜开关进入设计判断，系统按场景选择多姿态参考，并将参考图与角色限制一起提交生成。", "IP 使用闭环：")

    doc.add_heading("七、当前内容规模", level=1)
    add_table(
        doc,
        ["内容", "当前规模", "说明"],
        [
            ["启用预设", "4 个 + 默认模式", "手绘扁平涂鸦、3D、极简扁平插画、实景商品"],
            ["预设图像素材", "约 210 张", "4 个启用预设目录中的图像文件合计"],
            ["素材库", "246 条", "包含预设素材和用户保存素材，可继续动态增长"],
            ["创意方法", "12 张", "覆盖场景重构、感官置换、图形表达、纹理变化等方法"],
            ["Good / Bad Case", "1 组完整案例", "冰雪温泉季，包含 Brief、正反案例和评审依据"],
            ["工作区成图", "211 个图像文件", "当前工作区 outputs 中留存的生成与处理结果"],
            ["拆分包", "21 组", "标题层 / 背景层相关拆分结果"],
        ],
        [1.4, 1.3, 3.8],
    )
    p = doc.add_paragraph()
    r = p.add_run("注：以上为 2026 年 7 月 24 日工作区统计，用于说明原型规模，不等同于线上业务产量。")
    set_font(r, size=8.5, color=MID_GRAY)

    doc.add_heading("八、当前价值", level=1)
    add_table(
        doc,
        ["价值方向", "具体体现"],
        [
            ["更容易开始", "用户只需提供少量核心信息，系统负责补全设计过程，而不是要求用户先学会写 Prompt。"],
            ["更接近设计工作方式", "创意发散、方案筛选、参考分工、美术预检被纳入链路，不再直接从一句话跳到成图。"],
            ["更可控", "每张参考图有明确职责；上传主体和固定图层有硬约束；缺失信息不自动补。"],
            ["更可复用", "搜索、素材、预设、案例、生成结果和拆分资产都能回到下一次生产。"],
            ["更便于协作", "业务能看到中间判断，设计师能理解依据并接着修改，团队能逐步沉淀自己的视觉知识。"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("九、当前边界与风险", level=1)
    add_bullet(doc, "图像模型仍具有概率性：复杂中文、细节装饰和整合版式的像素级复刻不能保证一次完全准确。")
    add_bullet(doc, "整合版式是当前重点优化项：已完成策略与真实图片传递，但字体风格、比例和细节装饰的稳定遵循仍需更多测试集和评价指标。")
    add_bullet(doc, "质量和速度需要取舍：默认快速模式减少模型节点；开启生成后评审和自动重试会增加时间与成本。")
    add_bullet(doc, "Pinterest / Behance 当前为公开搜索验证方案，生产化仍需处理接口合规、版权、限流和来源授权。")
    add_bullet(doc, "当前为本地工作台原型，正式上线仍需要账户权限、任务队列、并发控制、监控告警、成本统计和数据库治理。")

    doc.add_heading("十、下一阶段建议", level=1)
    add_table(
        doc,
        ["优先级", "工作", "目标"],
        [
            ["P0", "建立整合版式评测集：标题字形、字号比例、对齐、原有装饰保留、无关文字增生", "让“参考图真的有用”变成可量化结果"],
            ["P0", "扩充 Good/Bad Case 与创意方法卡，并按主题、人群、场景建立覆盖", "提升创意质量与稳定性"],
            ["P1", "记录每次参考候选、选择原因、最终结果和人工评价", "逐步把参考匹配从规则升级为可学习排序"],
            ["P1", "完善任务队列、超时、重试、成本与耗时监控", "稳定达到 3 分钟内的体验目标"],
            ["P2", "加强生成后编辑：文字替换、局部重绘、图层导出与设计工具衔接", "提高真正生产可用性"],
        ],
        [0.65, 3.45, 2.4],
    )

    doc.add_heading("建议汇报结尾", level=1)
    add_callout(
        doc,
        "核心判断",
        "这个项目目前最有价值的成果，不只是“能生成一张图”，而是已经把需求、创意、参考、生成和资产连接成一条可以继续被团队训练和扩充的工作流。下一阶段的重点，是用更系统的案例、标签和评测，把“看起来不错”推进到“稳定可用”。",
        PALE_TEAL,
    )

    doc.save(REPORT_PATH)


def build_speech() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(
        doc,
        "5 分钟演讲稿",
        "营销 KV 智能生图工具｜轻松口语版",
        "SPEAKER NOTES",
        "建议时长约 5 分钟。现场可以直接照读，也可以只看每段开头的提示词自由发挥。",
    )
    set_repeatable_header_footer(doc, "营销 KV 智能生图工具｜5 分钟演讲稿")

    sections = [
        (
            "0:00–0:40｜先说为什么做",
            "大家好，我今天想介绍的是我们做的营销 KV 智能生图工具。\n\n"
            "做这个工具的原因其实很直接。平时业务或运营提设计需求，经常就是一句话，比如“帮我做一张冰雪温泉季的竖版海报”。这句话没有错，但真正做成一张好看的 KV，中间还缺很多判断：主题到底怎么表现，主体是谁，标题放哪，画面是冷还是暖，参考图应该参考字体还是构图，以及哪些东西绝对不能乱加。\n\n"
            "如果只是把这句话直接丢给生图模型，偶尔会出一张不错的，但更多时候靠运气。我们想解决的，就是怎么把这段原来藏在设计师脑子里的过程，变成一套系统可以执行的步骤。",
        ),
        (
            "0:40–2:40｜这套链路怎么工作",
            "现在用户输入标题、副标题、时间、画面描述和比例，也可以上传产品图、人物图，或者选择一个风格预设。\n\n"
            "第一步不是马上出图，而是先做 Brief 理解。这里尽量克制，只确认用户到底说了什么，哪些是硬要求，哪些只是可以推断，避免模型自己编品牌、卖点、英文或者价格。\n\n"
            "第二步是创意判断。系统会参考我们整理的创意方法和正反案例，先想三个不同方向，再自己淘汰两个，选一个最适合主题和画幅的。还是拿冰雪温泉季举例：只画滑雪，温泉没了；左右一半滑雪一半温泉，竖版很割裂；最后选的是一个完整雪景，把滑雪、温泉和远景小镇放进同一个画面。这样它不是一句“做得有创意”，而是已经有一张分层草图了。\n\n"
            "第三步是选参考图。以前很多工具只是随机抽一张，或者在 Prompt 里写“请参考这张图”。我们现在会先判断需要什么证据：哪张控制整套文字和版式，哪张控制风格质感，哪张控制元素造型，哪张控制角色动作。每张图都要说清楚为什么选、参考什么、什么不能抄。\n\n"
            "尤其是现在的整合版式，它不只管字体，还同时管主标题、副标题、时间的大小关系、对齐方式，以及文字区和画面区怎么避让。参考图里本来有引号、框线、标签或装饰，就保留；没有的就不乱加。主标题、副标题和时间再替换成用户自己的内容。\n\n"
            "最后，系统会形成设计大纲，做一次美术预检，再拼成最终 Prompt。最关键的一点是，最终生成时不只传 Prompt，选中的参考图和用户上传图也会真的一起传给 gpt-image-2。",
        ),
        (
            "2:40–4:20｜现在已经做到了什么",
            "目前生成页面已经可以完成完整流程，而且 Brief、设计判断和最终 Prompt 会一段一段出来，用户不用一直盯着空白页面等。\n\n"
            "风格方面，现在有默认模式和四个启用预设，包括手绘扁平、3D、极简扁平插画和实景商品。新的风格也不需要每次改代码，可以按文件夹结构上传，配置后直接在首页使用。\n\n"
            "素材这块也已经接起来了。用户可以自己上传、用表格导入，也可以从 Pinterest 和 Behance 搜索设计灵感。搜索到喜欢的图，可以保存到素材库，再点“用作参考图”回到生成页；也可以点“做同款”，把它的描述直接带进输入框。\n\n"
            "生成完的图会进入资产页，最新的排在最上面。用户可以删除、下载，还可以让 AI 把标题和背景拆成两张图，方便设计师继续处理。\n\n"
            "另外我们也做了几个更具体的控制。比如用户上传一瓶洗衣液并写“产品为图 1”，它必须保留这瓶产品的结构和包装，不能莫名其妙换成手机。兜兜 IP 也已经进入链路，系统会根据场景匹配合适姿态，而且角色必须出现并遵守固定造型限制。Logo 和右下角搜索框则是生成后的真实图层，用户可以选择开关，不让模型自己乱画。",
        ),
        (
            "4:20–5:00｜最后怎么总结",
            "所以目前我们已经跑通了几条比较完整的链路：从一句需求到成图，从外部灵感到素材库，从素材库再回到生成，从生成图到资产和拆分结果，以及从一个新风格文件夹到首页可用预设。\n\n"
            "我觉得这个项目现在最重要的成果，不只是“它能画图”，而是我们已经搭出了一套可以不断加案例、加方法、加参考图，并且越用越完善的工作流。\n\n"
            "当然它还不是终点。现在最需要继续优化的是整合版式和字体细节的稳定遵循，以及用更多 Good Case、Bad Case 做系统评测。我们的目标也不是完全替代设计师，而是让业务和运营先得到更像样的第一版，让设计师接手时不用从零开始。谢谢大家。",
        ),
    ]

    for title, text in sections:
        doc.add_heading(title, level=1)
        for paragraph_text in text.split("\n\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(paragraph_text)
            set_font(r, size=12)

    doc.add_heading("现场提醒", level=1)
    add_bullet(doc, "讲到“不是只传 Prompt”时可以停顿一下，这是当前方案与普通生图工具最明显的区别。")
    add_bullet(doc, "如果时间不够，可以删掉风格名称和素材规模，保留四条闭环。")
    add_bullet(doc, "如果领导追问效果稳定性，直接说明：主链路已跑通，整合版式的细节遵循仍是当前重点优化项。")
    doc.save(SPEECH_PATH)


if __name__ == "__main__":
    ROOT.mkdir(parents=True, exist_ok=True)
    build_report()
    build_speech()
    print(REPORT_PATH)
    print(SPEECH_PATH)
